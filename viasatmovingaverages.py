import pandas as pd
import datetime
import time
import seaborn as sns
import matplotlib.pyplot as plt

yesterday = (datetime.date.today() - datetime.timedelta(days = 1)).strftime('%m/%d/%Y')


ninety_prior = (datetime.date.today()-datetime.timedelta(days = 91)).strftime('%m/%d/%Y')

from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager

options = webdriver.ChromeOptions()
options.add_argument("--start-maximized")
prefs = {"profile.default_content_settings.popups": 0,
             "download.default_directory":
                        r"C:\Users\thegi\PycharmProjects\viasatmoving\\",
             "directory_upgrade": True}
options.add_experimental_option("prefs", prefs)

#driver = webdriver.Chrome(executable_path=r'C:\Users\thegi\PycharmProjects\viasatmoving\chromedriver.exe', options = options)
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
driver.maximize_window()
driver.get("https://fulfillment.wildblue.net")
time.sleep(5)

#x = input("user name")
#y = input("password")
#insert username and password below

elem = driver.find_element(By.NAME, 'j_username')
elem.send_keys("")
elem = driver.find_element(By.NAME,'j_password')
elem.send_keys("")
elem = driver.find_element(By.NAME,'submit')
elem.click()

driver.get("https://fulfillment.wildblue.net/fsm-fe/fsm/browseOrder/browseOrder.page?execution=e3s1")
try:
    elem = driver.find_element(By.ID, "browseOrder:orderGrid:filterBoard:createdDateFilter:j_idt172") or driver.find_element_by_id("browseOrder:orderGrid:filterBoard:createdDateFilter:j_idt196")
    elem.click()
except NoSuchElementException:
    elem = driver.find_element(By.ID,"browseOrder:orderGrid:filterBoard:createdDateFilter:j_idt189")
    elem.click()




elem = driver.find_element(By.ID, "browseOrder:orderGrid:filterBoard:scheduledDateFilter:dateFromCal_input")
elem.click()
elem.send_keys(ninety_prior)
elem = driver.find_element(By.ID, "browseOrder:orderGrid:filterBoard:scheduledDateFilter:dateToCal_input")
elem.click()
elem.send_keys(yesterday)
elem = driver.find_element(By.ID, 'browseOrder:orderGrid:search')
elem.click()
time.sleep(45)
elem = driver.find_element(By.NAME, 'browseOrder:xlsExport')
elem.click()
driver.minimize_window()
time.sleep(55)

import os

path = r'C:\Users\thegi\PycharmProjects\viasatmoving'

files = []
# r=root, d=directories, f = files
for r, d, f in os.walk(path):
    for file in f:
        if '.XLS' in file:
            files.append(os.path.join(r, file))

for f in files:
    print(f)

df = pd.read_excel(f)

fifteen_prior = (datetime.date.today()-datetime.timedelta(days = 16)).strftime('%m/%d/%Y')

fifteen_day_avg = (df[df['Completed date'].apply(lambda x: datetime.datetime.strptime(x[0:-13],'%m/%d/%Y'))>= fifteen_prior])

print('FIFTEEN DAY MONTHLY AVERAGE')
print(fifteen_day_avg['Order type'].value_counts()*2)
fifteen = (fifteen_day_avg['Order type'].value_counts()*2)

fifteen_vs1 = (fifteen_day_avg[fifteen_day_avg['Satellite Id'] == 'ViaSat-1'])


print('FIFTEEN DAY MONTHLY AVERAGE VIASAT 1')
print(fifteen_vs1['Order type'].value_counts()*2)
fifteen_visat1 = (fifteen_vs1['Order type'].value_counts()*2)


print('')

thirty_prior = (datetime.date.today()-datetime.timedelta(days = 31)).strftime('%m/%d/%Y')

thirty_day_avg = (df[df['Completed date'].apply(lambda x: datetime.datetime.strptime(x[0:-13],'%m/%d/%Y'))>= thirty_prior])

print('Thirty DAY MONTHLY AVERAGE')
print(thirty_day_avg['Order type'].value_counts())
thirty = (thirty_day_avg['Order type'].value_counts())

thirty_vs1 = (thirty_day_avg[thirty_day_avg['Satellite Id'] == 'ViaSat-1'])

print('Thirty DAY MONTHLY AVERAGE VIASAT 1')
print(thirty_vs1['Order type'].value_counts())
thirty_viasat1 = (thirty_vs1['Order type'].value_counts())

print('')

sixty_prior = (datetime.date.today()-datetime.timedelta(days = 61)).strftime('%m/%d/%Y')

sixty_day_avg = (df[df['Completed date'].apply(lambda x: datetime.datetime.strptime(x[0:-13],'%m/%d/%Y'))>= sixty_prior])

print('Sixty DAY MONTHLY AVERAGE')
print(sixty_day_avg['Order type'].value_counts()/2)
sixty = (sixty_day_avg['Order type'].value_counts()/2)

sixty_vs1 = (sixty_day_avg[sixty_day_avg['Satellite Id'] == 'ViaSat-1'])

print('Sixty DAY MONTHLY AVERAGE VIASAT 1')
print(sixty_vs1['Order type'].value_counts()/2)
sixty_viasat1 = (sixty_vs1['Order type'].value_counts()/2)

print('')

ninety_prior = (datetime.date.today()-datetime.timedelta(days = 91)).strftime('%m/%d/%Y')

ninety_day_avg = (df[df['Completed date'].apply(lambda x: datetime.datetime.strptime(x[0:-13],'%m/%d/%Y'))>= ninety_prior])

print('Ninety DAY MONTHLY AVERAGE')
print(ninety_day_avg['Order type'].value_counts()/3)
ninety = (ninety_day_avg['Order type'].value_counts()/3)

ninety_vs1 = (ninety_day_avg[ninety_day_avg['Satellite Id'] == 'ViaSat-1'])

print('Ninety DAY MONTHLY AVERAGE VIASAT 1')
print(ninety_vs1['Order type'].value_counts()/3)
ninety_viasat1 = (ninety_vs1['Order type'].value_counts()/3)

print('')

# PLOTS

sns.countplot(x ='Order type', data = fifteen_day_avg, palette = 'coolwarm', hue='Satellite Id')
plt.xticks(rotation=90)
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('fifteendayavg.png')
plt.clf()

sns.countplot(x ='Service region', data = fifteen_day_avg, palette = 'coolwarm', hue='Order type')
plt.xticks(rotation=90)
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('fifteendayavgservice.png')
plt.clf()

sns.countplot(x ='Technician last name', data = fifteen_day_avg, palette = 'coolwarm', hue='Order type')
plt.xticks(rotation=90)
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('fifteendayavgtech.png')
plt.clf()


sns.countplot(x ='Order type', data = thirty_day_avg, palette = 'coolwarm', hue='Satellite Id')
plt.xticks(rotation=90)
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('Thirtydayavg.png')
plt.clf()

sns.countplot(x ='Service region', data = thirty_day_avg, palette = 'coolwarm', hue='Order type')
plt.xticks(rotation=90)
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('thirtydayavgservice.png')
plt.clf()

sns.countplot(x ='Technician last name', data = thirty_day_avg, palette = 'coolwarm', hue='Order type')
plt.xticks(rotation=90)
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('thirtydayavgtech.png')
plt.clf()

sns.countplot(x ='Order type', data = sixty_day_avg, palette = 'coolwarm', hue='Satellite Id')
plt.xticks(rotation=90)
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('sixtydayavg.png')
plt.clf()

sns.countplot(x ='Service region', data = sixty_day_avg, palette = 'coolwarm', hue='Order type')
plt.xticks(rotation=90)
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('sixtydayavgservice.png')
plt.clf()

sns.countplot(x ='Technician last name', data = sixty_day_avg, palette = 'coolwarm', hue='Order type')
plt.xticks(rotation=90)
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('sixtydayavgtech.png')
plt.clf()



sns.countplot(x ='Order type', data = ninety_day_avg, palette = 'coolwarm', hue='Satellite Id')
plt.xticks(rotation=90)
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('ninetydayavg.png')

sns.countplot(x ='Service region', data = ninety_day_avg, palette = 'coolwarm', hue='Order type')
plt.xticks(rotation=90)
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('ninetydayavgservice.png')
plt.clf()

sns.countplot(x ='Technician last name', data = ninety_day_avg, palette = 'coolwarm', hue='Order type')
plt.xticks(rotation=90)
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('ninetydayavgtech.png')
plt.clf()

df['Completed date'] = pd.to_datetime(df['Completed date'])
df['dayofweek'] = df['Completed date'].apply(lambda time: datetime.datetime.weekday(time))
dmap = {0: 'Mon', 1: 'Tue', 2: 'Wed', 3: 'Thu', 4: 'Fri', 5: 'Sat', 6: 'Sun'}
df['dayofweek'] = df['dayofweek'].map(dmap)
sns.countplot(x = 'dayofweek', data = df, hue = 'Order type', palette = 'coolwarm')
plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='x-small')
plt.tight_layout()

plt.savefig('dayofweek.png')
plt.clf()



import docx
from docx import Document
from docx.shared import Inches

document = Document()

document.add_picture('viasat.png', width=Inches(1))

document.add_heading('Viasat Moving Monthly Averages', 0)

document.add_paragraph(f'FIFTEEN DAY')
document.add_picture('fifteendayavg.png', width=Inches(7))
document.add_paragraph(f'{fifteen}')
document.add_paragraph(f'FIFTEEN DAY VS1')
document.add_paragraph(f'{fifteen_visat1}')
document.add_paragraph(f'FIFTEEN DAY BY SERVICE REGION')
document.add_picture('fifteendayavgservice.png', width=Inches(7))
document.add_paragraph(f'FIFTEEN DAY BY TECH')
document.add_picture('fifteendayavgtech.png', width=Inches(7))
document.add_page_break()
document.add_paragraph(f'THIRTY DAY')
document.add_picture('Thirtydayavg.png', width=Inches(7))
document.add_paragraph(f'{thirty}')
document.add_paragraph(f'THIRTY DAY VS1')
document.add_paragraph(f'{thirty_viasat1}')
document.add_paragraph(f'THIRTY DAY BY SERVICE REGION')
document.add_picture('thirtydayavgservice.png', width=Inches(7))
document.add_paragraph(f'THIRTY DAY BY TECH')
document.add_picture('thirtydayavgtech.png', width=Inches(7))
document.add_page_break()
document.add_paragraph(f'SIXTY DAY')
document.add_picture('sixtydayavg.png', width=Inches(7))
document.add_paragraph(f'{sixty}')
document.add_paragraph(f'SIXTY DAY VS1')
document.add_paragraph(f'{sixty_viasat1}')
document.add_paragraph(f'SIXTY DAY BY SERVICE REGION')
document.add_picture('sixtydayavgservice.png', width=Inches(7))
document.add_paragraph(f'SIXTY DAY BY TECH')
document.add_picture('sixtydayavgtech.png', width=Inches(7))
document.add_page_break()
document.add_paragraph(f'NINETY DAY')
document.add_picture('ninetydayavg.png', width=Inches(7))
document.add_paragraph(f'{ninety}')
document.add_paragraph(f'NINETY DAY VS1')
document.add_paragraph(f'{ninety_viasat1}')
document.add_paragraph(f'NINETY DAY BY SERVICE REGION')
document.add_picture('ninetydayavgservice.png', width=Inches(7))
document.add_paragraph(f'NINETY DAY BY TECH')
document.add_picture('ninetydayavgtech.png', width=Inches(7))
document.add_page_break()
document.add_picture('dayofweek.png', width=Inches(7))



document.save(f'ViasatMovingAverages.docx')

os.remove(f)
os.system(f'start ViasatMovingAverages.docx')
